"""Output generation blocks for the block-based agentic pipeline system."""

import io
import os
from typing import Any, Dict, Optional, Set

import markdown
from docx import Document

from block_agents.core.block import Block
from block_agents.core.context import Context
from block_agents.core.errors import InputValidationError, OutputValidationError
from block_agents.core.registry import register_block


@register_block("output_generator")
class OutputGeneratorBlock(Block):
    """Block for generating formatted output.

    This block allows users to format the output of a pipeline in various formats:
    - Markdown
    - HTML
    - DOCX
    - Plain text
    """

    def __init__(self, block_id: str, config: Dict[str, Any]):
        """Initialize a new OutputGeneratorBlock.

        Args:
            block_id: Unique identifier for this block instance
            config: Block configuration
        """
        super().__init__(block_id, config)
        
        # Get format
        self.format = config.get("format", "markdown")  # markdown, html, docx, text
        
        # Get output options
        self.title = config.get("title", "")
        self.css = config.get("css", "")
        self.template = config.get("template", "")
        self.save_path = config.get("save_path", "")

    def process(self, inputs: Dict[str, Any], context: Context) -> Dict[str, Any]:
        """Process the inputs and produce an output.

        Args:
            inputs: Input values for the block
            context: Execution context

        Returns:
            Dictionary containing the formatted output
        """
        # Get text from inputs
        text = inputs.get("text", "")
        
        # Format the output
        if self.format == "markdown":
            # Pass through as markdown
            formatted_output = text
            output_format = "markdown"
            
        elif self.format == "html":
            # Convert markdown to HTML
            html = markdown.markdown(text)
            
            # Add CSS if provided
            if self.css:
                html = f"<style>{self.css}</style>\n{html}"
                
            # Add title if provided
            if self.title:
                html = f"<h1>{self.title}</h1>\n{html}"
                
            formatted_output = html
            output_format = "html"
            
        elif self.format == "docx":
            # Create DOCX document
            doc = Document()
            
            # Add title if provided
            if self.title:
                doc.add_heading(self.title, level=1)
                
            # Add content
            doc.add_paragraph(text)
            
            # Save to in-memory buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Save to file if path provided
            if self.save_path:
                # Create directory if needed
                os.makedirs(os.path.dirname(os.path.abspath(self.save_path)), exist_ok=True)
                
                # Save the file
                with open(self.save_path, "wb") as f:
                    f.write(buffer.getvalue())
                    
                context.log(self.id, f"Saved DOCX to {self.save_path}")
                
            formatted_output = buffer.getvalue()
            output_format = "docx"
            
        else:  # text
            # Format as plain text
            formatted_output = text
            output_format = "text"
            
        # Log the output format
        context.log(self.id, f"Generated output in {output_format} format")
        
        return {
            "output": formatted_output,
            "format": output_format,
            "title": self.title,
            "save_path": self.save_path if self.save_path else None,
        }

    def validate_inputs(self, inputs: Dict[str, Any]) -> None:
        """Validate the inputs to the block.

        Args:
            inputs: Input values for the block

        Raises:
            InputValidationError: If validation fails
        """
        if "text" not in inputs:
            raise InputValidationError(
                "Required input 'text' not found",
                block_id=self.id,
            )
            
        if not isinstance(inputs["text"], str):
            raise InputValidationError(
                "Input 'text' must be a string",
                block_id=self.id,
                details={"input_type": type(inputs["text"]).__name__},
            )

    def validate_output(self, output: Any) -> None:
        """Validate the output of the block.

        Args:
            output: Block output

        Raises:
            OutputValidationError: If validation fails
        """
        if not isinstance(output, dict):
            raise OutputValidationError(
                "Output must be a dictionary",
                block_id=self.id,
                details={"output_type": type(output).__name__},
            )
            
        if "output" not in output:
            raise OutputValidationError(
                "Output dictionary must contain 'output' key",
                block_id=self.id,
            )
            
        if "format" not in output:
            raise OutputValidationError(
                "Output dictionary must contain 'format' key",
                block_id=self.id,
            )

    def get_required_inputs(self) -> Set[str]:
        """Get the set of required input keys for this block.

        Returns:
            Set of required input keys
        """
        return {"text"}

    def get_optional_inputs(self) -> Set[str]:
        """Get the set of optional input keys for this block.

        Returns:
            Set of optional input keys
        """
        return set()